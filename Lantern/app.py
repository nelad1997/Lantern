import streamlit as st
import base64
import time
from streamlit_quill import st_quill
import re
import json
import os
import fitz  # PyMuPDF
from docx import Document
import io
import html  # Added for safe text escaping
import sys
import PIL.Image
import logging
import datetime

# --- On-Screen Debug Helper ---
def add_debug_log(msg):
    if "debug_logs" not in st.session_state:
        st.session_state.debug_logs = []
    timestamp = datetime.datetime.now().strftime("%H:%M:%S")
    st.session_state.debug_logs.append(f"[{timestamp}] {msg}")
    # Also log to terminal
    logging.getLogger(__name__).info(msg)

# Set up logging for app.py
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# --- Environment Hardening for Streamlit Cloud ---
# Explicitly add the script's directory to sys.path to ensure local imports like 'definitions' are reliable
current_dir = os.path.dirname(os.path.abspath(__file__))
if current_dir not in sys.path:
    sys.path.append(current_dir)

# --- Imports ---
from definitions import UserEventType, ActionType
from tree import init_tree, get_current_node, navigate_to_node, get_node_short_label, get_nearest_html, load_tree, save_tree
from controller import handle_event, generate_diff_html, apply_fuzzy_replacement
from sidebar_map import render_sidebar_map
from dotenv import load_dotenv

# --- ReportLab Import Handling (Safe Import) ---
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    from reportlab.lib.utils import simpleSplit

    HAS_REPORTLAB = True
except ImportError:
    HAS_REPORTLAB = False

load_dotenv()

# -------------------------------------------------
# Page Configuration & Logo Loading
# -------------------------------------------------
# Robust Path Hunting: script dir, then cwd
logo_search_paths = [
    os.path.join(current_dir, "logo.jpg"),
    os.path.join(os.getcwd(), "logo.jpg")
]
logo_full_path = None
for p in logo_search_paths:
    if os.path.exists(p):
        logo_full_path = p
        break

page_logo = "🏮"
logo_base64 = None

if logo_full_path:
    try:
        # Load for Page Icon
        page_logo = PIL.Image.open(logo_full_path)
        # Load for base64 Markdown display
        with open(logo_full_path, "rb") as f:
            logo_base64 = base64.b64encode(f.read()).decode()
        logger.info(f"✅ LOGO LOAD SUCCESS: {logo_full_path} (base64 length: {len(logo_base64) if logo_base64 else 0})")
    except Exception as e:
        logger.error(f"❌ LOGO LOAD ERROR: {e}")
        page_logo = "🏮"

st.set_page_config(page_title="Lantern", page_icon=page_logo, layout="wide", initial_sidebar_state="collapsed")


# -------------------------------------------------
# File Processing Helpers
# -------------------------------------------------
def extract_text_from_file(uploaded_file):
    # Reset file pointer to beginning (Crucial for Streamlit)
    uploaded_file.seek(0)

    file_type = uploaded_file.name.split('.')[-1].lower()

    try:
        if file_type == 'pdf':
            doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
            text = ""
            for page in doc:
                text += page.get_text()
            return text

        elif file_type == 'docx':
            doc = Document(io.BytesIO(uploaded_file.read()))
            return "\n".join([para.text for para in doc.paragraphs])

        elif file_type in ['txt', 'md']:
            return uploaded_file.getvalue().decode("utf-8")

    except Exception as e:
        st.error(f"Error reading file: {e}")
        return None

    return None


# -------------------------------------------------
# Export Helpers
# -------------------------------------------------
def create_docx(text):
    doc = Document()
    for paragraph in text.split('\n'):
        if paragraph.strip():
            doc.add_paragraph(paragraph)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def create_pdf(text):
    if not HAS_REPORTLAB:
        return None

    bio = io.BytesIO()
    c = canvas.Canvas(bio, pagesize=letter)
    width, height = letter

    c.setFont("Helvetica", 12)

    y = height - 50
    margin = 50

    for line in text.split('\n'):
        wrapped_lines = simpleSplit(line, "Helvetica", 12, width - 2 * margin)
        for wrapped_line in wrapped_lines:
            if y < 50:
                c.showPage()
                c.setFont("Helvetica", 12)
                y = height - 50
            c.drawString(margin, y, wrapped_line)
            y -= 15
        y -= 5

    c.save()
    return bio.getvalue()


# -------------------------------------------------
# Persistence Helpers
# -------------------------------------------------
AUTOSAVE_FILE = "lantern_autosave.json"


def save_autosave(tree):
    """
    DISABLED: This function intentionally does nothing.
    No file will be created on the disk.
    """
    pass


def load_autosave():
    """Tries to load the tree state from the local autosave file (if exists)."""
    if not os.path.exists(AUTOSAVE_FILE):
        return False
    try:
        with open(AUTOSAVE_FILE, 'r', encoding='utf-8') as f:
            data = json.load(f)
        st.session_state.tree = data
        if "pinned_items" not in st.session_state.tree:
            st.session_state.tree["pinned_items"] = []
        return True
    except Exception as e:
        return False


def save_project(tree):
    """Exports the project to a JSON string (in memory only)."""
    return json.dumps(tree, indent=2, ensure_ascii=False)


def load_project(json_str):
    try:
        data = json.loads(json_str)
        st.session_state.tree = data
        if "pinned_items" not in st.session_state.tree:
            st.session_state.tree["pinned_items"] = []
        st.session_state.banned_ideas = []
        # save_autosave is disabled
        return True
    except Exception as e:
        st.error(f"Failed to load project: {e}")
        return False


# -------------------------------------------------
# Image / Assets
# -------------------------------------------------
# LOGO Loading moved to the top of the script for global availability and robustness.

# -------------------------------------------------
# Styling (CSS)
# -------------------------------------------------
st.markdown("""
<style>
.stButton button { border-radius: 8px; font-weight: 500; transition: all 0.2s; }
.sidebar-header { display: flex; flex-direction: column; align-items: center; margin-bottom: 5px; margin-top: -50px; }
.sidebar-logo { max-width: 80px; width: 80px; height: auto; opacity: 0.8; margin-bottom: 0px; filter: grayscale(10%); }
.pinned-box { background-color: #fefce8; padding: 12px; border-radius: 6px; border-left: 4px solid #eab308; margin-bottom: 8px; font-size: 0.85rem; color: #422006; }
.suggestion-card { background-color: white; border: 1px solid #e2e8f0; border-radius: 8px; padding: 12px; margin-bottom: 12px; box-shadow: 0 1px 2px rgba(0,0,0,0.05); }
.suggestion-meta { display: flex; justify-content: space-between; font-size: 0.75rem; color: #94a3b8; margin-bottom: 8px; text-transform: uppercase; }
.suggestion-text { font-size: 0.95rem; color: #334155; line-height: 1.5; margin-bottom: 12px; }
.status-pill { display: inline-block; padding: 10px 24px; border-radius: 35px; font-size: 1.1rem; font-weight: 700; text-transform: uppercase; letter-spacing: 0.5px; box-shadow: 0 4px 6px -1px rgba(0,0,0,0.1), 0 2px 4px -1px rgba(0,0,0,0.06); width: 100%; text-align: center; }
.status-explore { background-color: #e0f2fe; color: #0369a1; border: 2px solid #bae6fd; }
.status-reflect { background-color: #fef3c7; color: #92400e; border: 2px solid #fde68a; }
.status-ready { background-color: #f0fdf4; color: #166534; border: 2px solid #bbf7d0; }
.status-structure { background-color: #ecfeff; color: #0891b2; border: 2px solid #a5f3fc; }
.status-refining { background-color: #f5f3ff; color: #5b21b6; border: 2px solid #ddd6fe; }
.action-bar.status-explore { background-color: #e0f2fe; border-color: #bae6fd; }
.action-bar.status-reflect { background-color: #fef3c7; border-color: #fde68a; }
.action-bar.status-refining { background-color: #f5f3ff; border-color: #ddd6fe; }
.action-bar.status-structure { background-color: #ecfeff; border-color: #a5f3fc; }
.action-bar {
    border: 1px solid #e5e7eb;
    border-radius: 14px;
    padding: 14px 16px 16px 16px;
    margin-bottom: 16px;
    background-color: #f9fafb;
    transition: background-color 0.3s ease, border-color 0.3s ease;
}
.action-bar.status-explore { background-color: #f0f9ff; border-color: #bae6fd; }
.action-bar.status-reflect { background-color: #fffbeb; border-color: #fde68a; }
.action-bar.status-refining { background-color: #f5f3ff; border-color: #ddd6fe; }
.action-bar.status-structure { background-color: #f0f9ff; border-color: #bae6fd; }
.action-bar-title {
    font-size: 0.9rem;
    font-weight: 600;
    color: #475569;
    margin-bottom: 10px;
}
/* Enhanced Typography */
.ql-editor {
    font-size: 18px !important;
    line-height: 1.6 !important;
}
.ql-container {
    font-size: 18px !important;
}
/* Consistent Scrollable Content for all panels */
/* Standard scrollable content boxes (Original/Proposed/Critique/Expand) */
.scrollable-content {
    max-height: 250px; 
    overflow-y: auto !important;
    overflow-x: hidden;
    display: block !important;
    width: 100%;
    margin-bottom: 5px; 
    background-color: #f8fafc;
    padding: 10px;
    border-radius: 6px;
    border: 1px solid #e2e8f0;
    font-size: 0.88rem;
    color: #334155;
    line-height: 1.5;
    white-space: pre-wrap;
}
.scrollable-content::-webkit-scrollbar {
    width: 4px;
}
.scrollable-content::-webkit-scrollbar-thumb {
    background-color: #cbd5e1;
    border-radius: 2px;
}
/* Specific tree scroll container - ChatGPT fix */
.tree-scroll-box {
    max-height: 55vh;
    overflow-y: auto;
    overflow-x: hidden;
    padding-right: 6px;
    border: 1px solid #e5e7eb;
    border-radius: 8px;
}
.tree-scroll-box svg {
    height: auto !important;
    max-width: 100%;
}

/* New IMG-based scrolling solution */
.tree-scroll-img {
    max-height: 55vh;
    overflow: auto;
    border: 1px solid #e5e7eb;
    border-radius: 8px;
    background: white;
}
.tree-scroll-img img {
    display: block;
    width: auto !important;
    max-width: none !important; 
}
</style>
""", unsafe_allow_html=True)


# -------------------------------------------------
# Helpers
# -------------------------------------------------
def get_ui_state(tree):
    # 1. Active thinking or pending AI task takes priority
    if st.session_state.get("is_thinking", False):
         if st.session_state.pending_action:
             act = st.session_state.pending_action.get("action")
             if act == ActionType.DIVERGE: return "Expanding...", "status-explore"
             if act == ActionType.CRITIQUE: return "Critiquing...", "status-reflect"
             if act == ActionType.REFINE: return "Refining...", "status-refining"
             if act == ActionType.SEGMENT: return "Analyzing structure...", "status-structure"
         return "Thinking...", "status-explore"

    # 2. Review modes (if we have suggestions on screen) - Changed to Drafting as per user request
    if "last_refine_diff" in st.session_state or st.session_state.get("pending_refine_edits"):
        return "Drafting", "status-ready"
    
    if "current_critiques" in st.session_state and st.session_state["current_critiques"]:
        return "Drafting", "status-ready"

    # 3. Default to Drafting for everything else
    return "Drafting", "status-ready"


# -------------------------------------------------
# State Synchronization Callbacks
# -------------------------------------------------
def sync_focus_mode():
    choice = st.session_state.get("promo_focus_mode_radio", "Whole Document")
    st.session_state["promo_focus_mode"] = "Whole Document" if choice == "Whole Document" else "Specific Paragraph"

def sync_paragraph_selection():
    # Attempt to pull from the specific widget key
    selection = st.session_state.get("promo_block_radio_selector")
    if selection:
        try:
            # Flexible regex to catch [N], [P N], or just N at the start
            match = re.search(r"(\d+)", selection)
            if match:
                idx = int(match.group(1)) - 1
                st.session_state["promo_block_selector_idx"] = idx
        except:
            pass
    
    # Fallback to ensure we have a valid index if mode is Specific Paragraph
    if st.session_state.get("promo_focus_mode") == "Specific Paragraph":
        saved_idx = st.session_state.get("promo_block_selector_idx", 0)
        paras = st.session_state.get("structural_segments", [])
        if paras:
            st.session_state["promo_block_selector_idx"] = max(0, min(saved_idx, len(paras)-1))

def get_document_structure(current_html):
    """
    Deterministic structural segmentation of HTML.
    Identifies units based on block-level tags and line breaks.
    Returns (title_string or None, list_of_paragraph_strings).
    """
    if not current_html or not current_html.strip():
        return None, []

    # Clean styling and irrelevant tags
    no_css = re.sub(r"<style.*?>.*?</style>", "", current_html, flags=re.DOTALL | re.IGNORECASE)
    
    # 1. Pre-process common separators to preserve line breaks before stripping tags
    processed_html = no_css.replace("<br>", "\n").replace("<br/>", "\n").replace("<br />", "\n")
    processed_html = re.sub(r"</(p|div|h[1-6]|li|blockquote)>", r"</\1>\n\n", processed_html)

    # findall for block-level contents
    blocks = re.findall(r"<(p|h[1-6]|div|li|blockquote)[^>]*>(.*?)</\1>", processed_html, re.DOTALL | re.IGNORECASE)
    
    all_segments = []
    if blocks:
        for tag, content in blocks:
            clean_text = re.sub("<[^<]+?>", "", content).replace("&nbsp;", " ").strip()
            if not clean_text:
                continue
            raw_splits = [s.strip() for s in re.split(r"\n+", clean_text) if s.strip()]
            for ss in raw_splits:
                if len(ss) >= 3:
                    all_segments.append(ss)
    else:
        # Fallback for plain text or unstructured HTML
        plain_text = re.sub("<[^<]+?>", "", processed_html).replace("&nbsp;", " ").strip()
        all_segments = [p.strip() for p in re.split(r"\n+", plain_text) if len(p.strip()) >= 3]
            
    if not all_segments:
        return None, []

    # Heuristic: Identify if the first line is a Title
    first = all_segments[0]
    is_title_length = len(first) < 250
    ends_with_punct = first.endswith(('.', '?', '!', ':', ';'))
    
    # If first line is short and doesn't end in typical paragraph punctuation
    if is_title_length and not (len(first) > 100 and ends_with_punct):
        return first, all_segments[1:]
    
    return None, all_segments

# -------------------------------------------------
# Main App
# -------------------------------------------------
def main():
    if "pending_action" not in st.session_state:
        st.session_state.pending_action = None
    if "is_thinking" not in st.session_state:
        st.session_state.is_thinking = False
    if "root_topic_resolved" not in st.session_state:
        st.session_state.root_topic_resolved = False
    if "llm_in_flight" not in st.session_state:
        st.session_state.llm_in_flight = False
    if "tree" not in st.session_state:
        # EPHEMERAL MODE: Always start fresh on reload (User Request)
        # We explicitly do NOT load from disk.
        import uuid
        # Generate a random ID for this run so saves don't overwrite previous runs (even though we don't reload them)
        if "stable_session_id" not in st.session_state:
            st.session_state["stable_session_id"] = str(uuid.uuid4())[:8]

        st.session_state.tree = init_tree("")
        add_debug_log(f"✨ Started new ephemeral session: {st.session_state['stable_session_id']}")
    
    # DEFENSIVE & MIGRATION: Ensure all tree fields exist
    tree = st.session_state.tree
    if "pinned_items" not in tree: tree["pinned_items"] = []
    if "banned_ideas" not in tree: tree["banned_ideas"] = []
    if "dismissed_suggestions" not in tree: tree["dismissed_suggestions"] = set()
    if "bulletproof_history" not in tree: tree["bulletproof_history"] = set()
    if "current_critiques" not in tree: tree["current_critiques"] = []
    if "pending_refine_edits" not in tree: tree["pending_refine_edits"] = []

    # UI Backwards Compatibility (aliasing for readability in existing code)
    st.session_state.banned_ideas = tree["banned_ideas"]
    st.session_state.dismissed_suggestions = tree["dismissed_suggestions"]
    st.session_state.bulletproof_history = tree["bulletproof_history"]
    st.session_state.current_critiques = tree["current_critiques"]
    st.session_state.pending_refine_edits = tree["pending_refine_edits"]

    if "selected_paths" not in st.session_state:
        st.session_state.selected_paths = []
    if "knowledge_base" not in st.session_state:
        st.session_state.knowledge_base = {}
    if "pending_refine_edits" not in st.session_state:
        st.session_state.pending_refine_edits = []
    if "ai_info_message" not in st.session_state:
        st.session_state.ai_info_message = None
    if "just_applied_refine" not in st.session_state:
        st.session_state.just_applied_refine = False
    if "promo_block_selector_idx" not in st.session_state:
        st.session_state.promo_block_selector_idx = 0
    if "promo_focus_mode" not in st.session_state:
        st.session_state.promo_focus_mode = "Whole Document"
    if "debug_logs" not in st.session_state:
        st.session_state.debug_logs = []

    # --- CLOUD GUARDRAIL: Session State Integrity ---
    from definitions import IS_CLOUD
    if IS_CLOUD:
        # In Streamlit Cloud, session state can be fragile. 
        # We ensure critical keys exist but avoid destructive resets if possible.
        if "tree" not in st.session_state:
            add_debug_log("☁️ CLOUD STATE RECOVERY: Tree mission. Initializing new.")
            st.session_state.tree = init_tree("")
        if "editor_html" not in st.session_state:
            add_debug_log("☁️ CLOUD STATE RECOVERY: editor_html missing. Resetting to root content.")
            st.session_state.editor_html = get_nearest_html(st.session_state.tree, st.session_state.tree["current"])
        if "banned_ideas" not in st.session_state:
            st.session_state.banned_ideas = []
        if "pending_refine_edits" not in st.session_state:
            st.session_state.pending_refine_edits = []

    
    # --- GLOBAL STRUCTURAL SYNC ---
    # Ensure segments exist before ANY logic or UI starts
    current_html = st.session_state.get("editor_html", "")
    if current_html.strip():
        # Only re-fetch if empty to provide a stable substrate
        if not st.session_state.get("structural_segments"):
            _, paras = get_document_structure(current_html)
            st.session_state.structural_segments = paras
    else:
        st.session_state.structural_segments = []

    tree = st.session_state.tree
    current_node = get_current_node(tree)

    if "editor_html" not in st.session_state:
        st.session_state["editor_html"] = get_nearest_html(tree, st.session_state.tree["current"])
    
    # --- IMMEDIATE WIDGET SYNC (Fixes AI Focus Mismatch) ---
    # Sync focus widgets to state BEFORE calculating target_text/preview
    focus_choice = st.session_state.get("promo_focus_mode_radio", "Whole Document")
    st.session_state["promo_focus_mode"] = "Whole Document" if focus_choice == "Whole Document" else "Specific Paragraph"

    selection = st.session_state.get("promo_block_radio_selector")
    if selection:
        match = re.search(r"(\d+)", selection)
        if match:
            st.session_state["promo_block_selector_idx"] = int(match.group(1)) - 1
            
    mode_label, mode_class = get_ui_state(tree)

    col_editor, col_lantern = st.columns([2, 1], gap="large")

    # ==========================================
    # LEFT COLUMN: EDITOR
    # ==========================================
    with col_editor:
        # Header (Cleaned up, removed Import/Export from here)
        if "comparison_data" in st.session_state:
            st.subheader("⚖️ Branch Comparison")
        else:
            st.subheader("Editor")

        if "comparison_data" in st.session_state:
            comp = st.session_state.comparison_data
            diff_html = generate_diff_html(comp['a']['summary'], comp['b']['summary'])
            st.markdown(
                f'<div style="background-color: white; padding: 20px; border-radius: 8px; border: 1px solid #e2e8f0; line-height: 1.6;">{diff_html}</div>',
                unsafe_allow_html=True)

        # Logic moved above buttons, UI moved below
        combined_help_info = (
            "🤖 AI Actions Guide:&#10;&#10;"
            "• Expand: Generates 3 academic perspectives to grow your ideas.&#10;"
            "• Critique: Analyzes for logical gaps or ethical issues.&#10;"
            "• Refine: Suggests writing improvements.&#10;&#10;"
            "⚙️ Focus Settings:&#10;&#10;"
            "Choose between 'Whole Document' or 'Specific Paragraph' for AI focus.&#10;&#10;"
            "⚠️ IMPORTANT: If you make significant changes to your draft, please go to 'Paragraph Segmentation' below and click 'Refresh' to ensure the AI's focus remains accurate."
        )

        paragraphs_only = st.session_state.get("structural_segments", [])
        focus_mode = st.session_state.get("promo_focus_mode", "Whole Document")

        # Target text calculation for AI context
        target_text = ""
        block_idx = 1
        
        if focus_mode == "Specific Paragraph" and paragraphs_only:
            # Consistent index reading from synchronized state
            block_idx_raw = st.session_state.get("promo_block_selector_idx", 0)
            # Bounds check
            block_idx_raw = max(0, min(block_idx_raw, len(paragraphs_only) - 1))
            block_idx = block_idx_raw + 1
            target_text = paragraphs_only[block_idx_raw]
        else:
            # Robust Full Doc Extraction
            current_html = st.session_state.get("editor_html", "")
            # 1. Remove styles
            no_css = re.sub(r"<style.*?>.*?</style>", "", current_html, flags=re.DOTALL | re.IGNORECASE)
            # 2. Add newlines for block tags to preserve structure
            txt_s = re.sub(r"<(p|div|h[1-6]|li|blockquote|br)[^>]*>", "\n", no_css)
            txt_s = re.sub(r"</(p|div|h[1-6]|li|blockquote)>", "\n", txt_s)
            # 3. Strip all other tags and entities
            target_text = re.sub("<[^<]+?>", "", txt_s).replace("&nbsp;", " ").strip()
            # 4. Collapse extra newlines
            target_text = re.sub(r"\n{3,}", "\n\n", target_text)
            
            # Absolute fallback: If empty but HTML isn't, just strip everything
            if not target_text and current_html.strip():
                 target_text = re.sub("<[^<]+?>", "", current_html).strip()

        st.session_state["focused_text"] = target_text
        st.session_state["focus_scope_label"] = "Whole Document" if focus_mode == "Whole Document" else f"Paragraph {block_idx}"

        # AI Action Bar (Color coded, no redundant text)
        st.markdown(
            f'<div class="action-bar {mode_class}" style="display: flex; align-items: center; gap: 8px;">'
            f'<div class="action-bar-title" style="margin-bottom: 0;">AI Reasoning Actions</div>'
            f'<span title="{combined_help_info}" style="cursor: pointer; background-color: #38bdf8; color: white; border-radius: 50%; width: 20px; height: 20px; display: inline-flex; align-items: center; justify-content: center; font-size: 0.8rem; font-weight: bold; line-height: 1;">i</span>'
            f'</div>',
            unsafe_allow_html=True
        )

        c1, c2, c3 = st.columns([1, 1, 1], gap="small")
        
        with c1:
            if st.button("🌱 Expand", use_container_width=True, help="Explore alternative reasoning paths and divergent perspectives based on your focus."):
                st.session_state.pending_action = {
                    "action": ActionType.DIVERGE, 
                    "anchor_id": tree["current"]
                }
                st.session_state.is_thinking = True
                st.rerun()
        with c2:
            if st.button("⚖️ Critique", use_container_width=True, help="Analyze your reasoning for potential biases, gaps, or logical fallacies."):
                st.session_state.pending_action = {
                    "action": ActionType.CRITIQUE, 
                    "anchor_id": tree["current"]
                }
                st.session_state.is_thinking = True
                st.rerun()
        with c3:
            if st.button("✨ Refine", use_container_width=True, help="Generate granular writing suggestions and draft improvements for the selected focus."):
                st.session_state.pending_action = {
                    "action": ActionType.REFINE, 
                    "anchor_id": tree["current"]
                }
                st.session_state.is_thinking = True
                st.rerun()
        st.markdown("</div>", unsafe_allow_html=True)

        # --- (3) AI Context & Structure "Folder" ---
        # --- (3) AI Context & Structure "Folder" ---
        with st.expander("🧠 AI Context & Structure", expanded=False):
            tab1, tab2, tab3 = st.tabs(["🎯 Focus Range", "📑 Segmentation", "👁️ Focus Preview"])
            
            with tab1:
                # Sub-section: Focus Range
                st.markdown(
                    f'<div style="display: flex; align-items: center; gap: 8px; margin-top: 10px; margin-bottom: 5px;">'
                    f'<b>AI Focus Range</b> '
                    f'<span title="Select whether the AI should analyze your Whole Document or focus on a Specific Paragraph for more granular feedback." style="cursor: help; background-color: #38bdf8; color: white; border-radius: 4px; padding: 1px 8px; font-size: 0.7em; font-weight: bold;">i</span>'
                    f'</div>',
                    unsafe_allow_html=True
                )
                
                # Use globally initialized segments
                current_paras = st.session_state.get("structural_segments", [])

                focus_choice = st.radio(
                    "Select AI Focus Range:",
                    ["Whole Document", "Specific Paragraph"],
                    key="promo_focus_mode_radio",
                    horizontal=True,
                    label_visibility="collapsed"
                )

                if focus_choice == "Specific Paragraph":
                    if current_paras:
                        st.markdown("<div style='margin-top: 5px;'></div>", unsafe_allow_html=True)
                        options = []
                        for i, p in enumerate(current_paras):
                            # Clean label for radio
                            p_clean = re.sub(r"^(?:\[P\s*\d+\]|Block\s*\d+:?|\d+[\.)]|[*•\-])\s*", "", p, flags=re.IGNORECASE).strip()
                            preview = (p_clean[:60] + "...") if len(p_clean) > 60 else p_clean
                            
                            # All segments in this list are now considered paragraphs
                            label_idx = f"{i+1}"
                            options.append(f"[{label_idx}] {preview}")
                        
                        saved_idx = st.session_state.get("promo_block_selector_idx", 0)
                        # Bounds check
                        safe_idx = max(0, min(saved_idx, len(options)-1))
                            
                        st.radio(
                            "Select Target Paragraph:",
                            options=options,
                            index=safe_idx,
                            key="promo_block_radio_selector",
                            label_visibility="collapsed"
                        )
                    else:
                        st.warning("No paragraphs detected. Type some text first.")

            with tab2:
                # Sub-section: Paragraph Segmentation
                st.markdown(
                    f'<div style="display: flex; align-items: center; gap: 8px; margin-top: 10px; margin-bottom: 5px;">'
                    f'<b>Paragraph Segmentation</b> '
                    f'<span title="View the document units identified by Lantern. AI actions (like Critique or Refine) use these structural markers to provide targeted improvements. Use Refresh if you make major structural changes." style="cursor: help; background-color: #38bdf8; color: white; border-radius: 4px; padding: 1px 8px; font-size: 0.7em; font-weight: bold;">i</span>'
                    f'</div>',
                    unsafe_allow_html=True
                )
                paras = st.session_state.get("structural_segments", [])
                if paras:
                    st.markdown('<div style="font-size: 0.8rem; color: #64748b; margin-bottom: 8px;">Units identified in your draft:</div>', unsafe_allow_html=True)
                    with st.container(height=250, border=False):
                        for i, p in enumerate(paras):
                            p_clean = re.sub(r"^(?:\[P\s*\d+\]|Block\s*\d+:?|\d+[\.)]|[*•\-])\s*", "", p, flags=re.IGNORECASE).strip()
                            st.markdown(
                                f'<div style="margin-bottom: 8px; padding: 6px; background: #f8fafc; border-radius: 4px; border-left: 2px solid #38bdf8; font-size: 0.8rem;">'
                                f'<b>[P{i+1}]</b> {p_clean[:120]}...</div>',
                                unsafe_allow_html=True
                            )
                
                has_content = bool(st.session_state.get("editor_html", "").strip())
                if has_content:
                    st.markdown("<div style='margin-top: 5px;'></div>", unsafe_allow_html=True)
                    if st.button("🔄 Refresh Structure", use_container_width=True, help="Update the segments based on current changes (excluding title)."):
                        st.session_state.pending_action = {"action": ActionType.SEGMENT}
                        st.session_state.is_thinking = True
                        st.rerun()

            with tab3:
                # Sub-section: AI Focus Preview
                st.markdown(
                    f'<div style="display: flex; align-items: center; gap: 8px; margin-top: 10px; margin-bottom: 5px;">'
                    f'<b>AI Focus Preview</b> '
                    f'<span title="This peek shows the exact text and markers Lantern will send to the AI. Use this to verify that the selection is correct before running an action." style="cursor: help; background-color: #38bdf8; color: white; border-radius: 4px; padding: 1px 8px; font-size: 0.7em; font-weight: bold;">i</span>'
                    f'</div>',
                    unsafe_allow_html=True
                )
                st.caption("Exact text Lantern will analyze:")
                st.markdown(
                    f'<div style="font-size: 0.85rem; color: #475569; background-color: #f8fafc; padding: 10px; border-radius: 6px; border: 1px solid #e2e8f0; height: 180px; overflow-y: auto; white-space: pre-wrap;">'
                    f'{st.session_state.get("focused_text", "No text detected.")}'
                    f'</div>',
                    unsafe_allow_html=True
                )

        if st.session_state.get("ai_info_message"):
            c_msg, c_msg_del = st.columns([0.9, 0.1])
            c_msg.info(st.session_state["ai_info_message"])
            if c_msg_del.button("✖", key="clear_ai_info", help="Clear this message"):
                st.session_state["ai_info_message"] = None
                st.rerun()

        st.markdown("<div style='margin-bottom: 20px'></div>", unsafe_allow_html=True)
        EDITOR_CSS = "<style>.ql-editor { font-size: 18px !important; line-height: 1.6; }</style>"

        html_content = st.session_state["editor_html"]
        
        # --- NEW: Granular Refine Progress Bar ---
        if st.session_state.get("pending_refine_edits"):
            pending_count = len([p for p in st.session_state.pending_refine_edits if p["status"] == "pending"])
            if pending_count > 0:
                st.warning(f"✨ Reviewing {pending_count} suggested improvements in the sidebar.")
        
        if True: # Editor is now always visible
            # שינוי: עטיפת ה-Editor במיכל עם גובה קבוע המאפשר גלילה (ללא border כדי למנוע שגיאות גרסה)
            with st.container(height=600):
                # Small Floating Clear Button inside container
                st.markdown(
                    """
                    <div style="position: relative; height: 0; z-index: 1000; text-align: right; top: 10px; right: 10px; pointer-events: none;">
                    </div>
                    """, unsafe_allow_html=True
                )
                
                # Use columns for right-alignment of the button
                c_btn_1, c_btn_2 = st.columns([0.94, 0.06])
                if c_btn_2.button("🗑", help="Reset All: Clears the editor and resets the entire Thought Tree and AI context."):
                    # FULL RESET: Wipe everything
                    st.session_state["editor_html"] = ""
                    st.session_state.tree = init_tree("")
                    
                    # Wiping session state context
                    st.session_state.banned_ideas = []
                    st.session_state.dismissed_suggestions = set()
                    st.session_state.bulletproof_history = set()
                    st.session_state.selected_paths = []
                    st.session_state.current_critiques = []
                    st.session_state.pending_refine_edits = []
                    st.session_state.structural_segments = []
                    st.session_state.logical_paragraphs = []
                    st.session_state.focused_text = ""
                    
                    # Force Quill Re-mount to show empty editor
                    if "editor_version" not in st.session_state:
                        st.session_state.editor_version = 0
                    st.session_state.editor_version = st.session_state.get("editor_version", 0) + 1
                    
                    from tree import save_tree
                    save_tree(st.session_state.tree)

                # Prepend CSS to the value so it renders inside the iframe
                quill_value = EDITOR_CSS + st.session_state["editor_html"]
                
                html_content = st_quill(
                    value=quill_value,
                    placeholder="Start drafting...",
                    html=True,
                    key="editor", # Rule 2: Constant key
                )

                if html_content is not None:
                    # Robust Clean: Remove any <style> blocks
                    clean_html = re.sub(r"<style.*?>.*?</style>", "", html_content, flags=re.DOTALL | re.IGNORECASE)
                    
                    # Sync to state
                    current_html_state = st.session_state.get("editor_html", "")
                    
                    if st.session_state.get("just_applied_refine"):
                        st.session_state.just_applied_refine = False
                    elif clean_html.strip() != current_html_state.strip():
                        # TEXT CHANGE DETECTED
                        st.session_state["editor_html"] = clean_html
                        
                        # Rule 6: No implicit saving to tree metadata. 
                        # Tree is updated ONLY on explicit navigation / action.
                        
                        # Update paragraphs from current clean_html for UI only
                        _, paragraphs = get_document_structure(clean_html)
                        st.session_state.structural_segments = paragraphs
                        
                        # Update paragraphs from current clean_html
                        doc_title, paragraphs = get_document_structure(clean_html)
                        st.session_state.structural_segments = paragraphs # Sync for ALL node types
                        
                        # Update Root Label if this is the root node
                        if current_node.get("type") == "root":
                            if doc_title:
                                title_topic = re.sub(r"^(?:\[P\s*\d+\]|Block\s*\d+:?|\d+[\.)]|[*•\-])\s*", "", doc_title, flags=re.IGNORECASE).strip()
                                if len(title_topic) > 60: title_topic = title_topic[:57] + "..."
                                current_node["metadata"]["label"] = f"[{title_topic}]"
                            elif paragraphs:
                                current_node["metadata"]["label"] = f"[{paragraphs[0][:25]}...]"
                        
                        st.session_state.last_edit_time = time.time()

                
                # --- Automatic Segmentation Disabled (Fix for Rate Limits) ---
                # Segmentation now only runs via "Refresh Logical Map" or after explicit AI actions if desired
                pass




    # ==========================================
    # SIDEBAR TOOLS
    # ==========================================
    with st.sidebar:
        st.markdown(
            """
            <style>
            .stRadio [data-testid="stWidgetLabel"] { display: none; }
            [data-testid="stSidebar"] hr { margin: 1rem 0; }
            </style>
            """, unsafe_allow_html=True
        )
        # 1. State Cleanup & View Management
        if st.session_state.get("pending_refine_edits"):
            # Do NOT auto-switch back to Map. Let the user decide when to go back.
            # But we can still cleanup the list if the user explicitly clicks a "Clear All" or "Finish"
            pass
        
        # 2. Sidebar View Selector
        
        # Determine default index logic:
        # If we have PENDING edits, default to Refine View (index 1) unless user manually switched.
        # If we just applied a refine, we WANT to stay in Refine View to see others.
        default_index = 0
        has_pending = st.session_state.get("pending_refine_edits") and any(p["status"] == "pending" for p in st.session_state.pending_refine_edits)
        
        if has_pending:
             default_index = 1
        
        # If 'sidebar_view_toggle' isn't in state yet, set it based on logic
        if "sidebar_view_toggle" not in st.session_state:
             st.session_state.sidebar_view_toggle = ["🗺️ Thought Map", "✨ Refine Review"][default_index]

        # Force view to Refine if we have pending items and just applied one (to prevent flipping back)
        if st.session_state.get("just_applied_refine") and has_pending:
             default_index = 1

        sidebar_view = st.radio(
            "View:",
            ["🗺️ Thought Map", "✨ Refine Review"],
            index=default_index,
            horizontal=True,
            key="sidebar_view_toggle",
            label_visibility="collapsed"
        )
        st.divider()

    if sidebar_view == "🗺️ Thought Map":
        render_sidebar_map(tree, show_header=False)
    else:
        # --- NEW: Granular Refine Review Panel (Sidebar Style) ---
        with st.sidebar:
            if st.session_state.get("pending_refine_edits"):
                # (Existing logic for showing items...)
                c_titles_1, c_titles_2 = st.columns([0.85, 0.15])
                c_titles_1.markdown("### ✨ Refine Review")
                if c_titles_2.button("🗑", help="Dismiss all pending suggestions", key="dismiss_all_refine"):
                    st.session_state.pending_refine_edits = []
                    st.rerun()
                
                # Show AI info message if any (e.g. "No improvements found")
                if st.session_state.get("ai_info_message"):
                    st.info(st.session_state["ai_info_message"])

                st.info("Review AI suggestions and apply them individually.")
                
                # Use a standard container (no fixed height) to allow sidebar scroll
                with st.container():
                    for i, proposal in enumerate(st.session_state.pending_refine_edits):
                        if proposal["status"] != "pending": continue
                        
                        # Suggestion Title with Info Icon and Focus Label
                        suggestion_type = proposal.get("type", "Improvement")
                        scope = proposal.get("scope", "Whole Document")
                        st.markdown(
                            f'<div style="display: flex; align-items: center; gap: 8px; margin-top: 15px; margin-bottom: 5px;">'
                            f'<b style="font-size: 1rem;">{suggestion_type} {i+1}</b>'
                            f'<span title="Type: {suggestion_type}&#10;Academic Reasoning:&#10;{proposal["reason"]}&#10;Focus: {scope}" style="cursor: help; background-color: #38bdf8; color: white; border-radius: 4px; padding: 1px 8px; font-size: 0.8em; font-weight: bold;">i</span>'
                            f'<span style="background-color: #f1f5f9; color: #64748b; font-size: 0.75rem; padding: 2px 6px; border-radius: 4px; margin-left: auto;">Focus: {scope}</span>'
                            f'</div>',
                            unsafe_allow_html=True
                        )

                        with st.container(border=True):
                            # Merged Diff View (Legacy style: Red/Green)
                            diff_html = generate_diff_html(proposal['original'], proposal['proposed'])
                            st.caption("Compare Changes")
                            st.markdown(
                                f'<div class="scrollable-content" style="max-height: 150px; overflow-y: auto; font-size: 0.95rem; background-color: #ffffff; border-color: #e2e8f0; margin-bottom: 12px; padding: 12px; border-radius: 6px; border: 1px solid; line-height: 1.6;">'
                                f'{diff_html}'
                                f'</div>', 
                                unsafe_allow_html=True
                            )
                            
                            c_app, c_dis = st.columns(2)
                            if c_app.button("✔ Apply", key=f"app_refine_{proposal['id']}", use_container_width=True):
                                # Replace in HTML
                                current_html = st.session_state["editor_html"]
                                refined_html = apply_fuzzy_replacement(current_html, proposal['original'], proposal['proposed'])
                                
                                if refined_html:
                                    st.session_state["editor_html"] = refined_html
                                    # Sync back to tree metadata for persistence
                                    st.session_state.tree["nodes"][st.session_state.tree["current"]].setdefault("metadata", {})["html"] = refined_html
                                    
                                    proposal["status"] = "applied"
                                    st.session_state.editor_version += 1
                                    st.session_state.just_applied_refine = True
                                    st.toast("✅ Applied suggested improvement!", icon="✨")
                                    st.rerun()
                                else:
                                    st.error("⚠️ Automated placement failed due to text/formatting mismatch.")
                                    with st.expander("🔍 Show Debug Info (Why it failed)"):
                                        st.write("**What the AI tried to replace:**")
                                        st.code(proposal['original'])
                                        st.write("**What the AI proposed:**")
                                        st.code(proposal['proposed'])
                                        st.info("The system couldn't find a matching segment in your draft precisely enough. You can copy the code above and paste it manually.")
                            
                            if c_dis.button("✖ Skip", key=f"dis_refine_{proposal['id']}", use_container_width=True):
                                proposal["status"] = "dismissed"
                                st.rerun()
                
                # --- After the loop: Completion state ---
                has_pending = any(p["status"] == "pending" for p in st.session_state.pending_refine_edits)
                if not has_pending:
                    st.success("✅ All suggestions have been reviewed!")
            else:
                # Fallback UI for empty Refine tab
                st.markdown("### ✨ Refine Review")
                if st.session_state.get("ai_info_message"):
                    st.info(st.session_state["ai_info_message"])
                else:
                    st.write("No active suggestions to review.")

    # ==========================================
    # RIGHT COLUMN: INTERACTION & CONTEXT
    # ==========================================
    with col_lantern:
        if logo_base64:
            st.markdown(
                f'<div style="display: flex; flex-direction: column; align-items: center; margin-bottom: 5px;"><img src="data:image/jpeg;base64,{logo_base64}" style="width: 70px; opacity: 0.9;"></div>',
                unsafe_allow_html=True)

        # --- NEW: Tooltip & System Explanation ---
        st.markdown(
            """
            <div style="text-align: center; margin-bottom: 15px;">
                <span title="Welcome to Lantern 💡&#10;What is Lantern?&#10;An intelligent environment for academic writing and reasoning. It combines a text editor with an AI assistant to facilitate in-depth research and thought management.&#10;How to use it?&#10;&#10;✍️ Write: Draft your initial ideas in the main text editor.&#10;🧠 Collaborate with AI:&#10;🌱 Expand: Explore new lines of reasoning and deepen your discussion.&#10;⚖️ Critique: Receive constructive feedback grounded in academic principles.&#10;✨ Refine: Polish your phrasing, precision, and style.&#10;🗺️ Navigate: Use the Thought Tree (sidebar) to manage various drafts and focus on specific sections of your work.&#10;&#10;⚠️ Pro-Tip: When refining your draft, periodically check the 'Paragraph Segmentation' section to ensure the building blocks are correctly mapped for the AI." style="cursor: help; color: #555; border-bottom: 1px dotted #777; font-size: 0.9em;">
                    ℹ️ How to use Lantern
                </span>
            </div>
            """,
            unsafe_allow_html=True
        )

        # --- NEW: Import / Export Buttons Moved Here ---
        c_io_1, c_io_2 = st.columns([1, 1])

        with c_io_1:
            import_tooltip = "Import a DOCX or PDF file to replace the current text in the editor."
            with st.popover("📥 Import", use_container_width=True, help=import_tooltip):
                # Dynamic key forces widget reset when editor_version changes (e.g. on Global Reset)
                uploader_key = f"doc_import_right_{st.session_state.get('editor_version', 0)}"
                uploaded_doc = st.file_uploader("Upload DOCX/PDF to replace content", type=["pdf", "docx", "txt", "md"],
                                                key=uploader_key)
                if uploaded_doc:
                    file_ext = uploaded_doc.name.split('.')[-1].lower()
                    if file_ext not in ['pdf', 'docx']:
                        st.error("Error: File type not supported. Please upload a DOCX or PDF file.")
                    else:
                        # Allow re-import if editor is empty OR if file changed
                        is_new_file = st.session_state.get("last_imported_doc") != uploaded_doc.name
                        is_empty_editor = not st.session_state.get("editor_html", "").strip()
                        
                        if is_new_file or is_empty_editor:
                            doc_text = extract_text_from_file(uploaded_doc)
                            if doc_text is not None:
                                lines = [l.strip() for l in doc_text.split('\n') if l.strip()]
                                html_blocks = []
                                for i, line in enumerate(lines):
                                    escaped = html.escape(line)
                                    if i == 0 and len(line) < 100:
                                        # Likely a title
                                        html_blocks.append(f"<h1>{escaped}</h1>")
                                    else:
                                        html_blocks.append(f"<p>{escaped}</p>")
                                
                                html_val = "\n".join(html_blocks)
                                st.session_state["editor_html"] = html_val
                                _, paragraphs = get_document_structure(html_val)
                                st.session_state.structural_segments = paragraphs # RE-SEGMENT on import (excluding title)
                                current_node.setdefault("metadata", {})["html"] = html_val
                                st.session_state.editor_version += 1
                                st.session_state["last_imported_doc"] = uploaded_doc.name
                                st.rerun()

        with c_io_2:
            export_tooltip = "Export your draft to DOCX or PDF format."
            with st.popover("📤 Export", use_container_width=True, help=export_tooltip):
                export_text = current_node.get("metadata", {}).get("draft_plain", "")
                if not export_text and "editor_html" in st.session_state:
                    export_text = re.sub("<[^<]+?>", "", st.session_state["editor_html"]).strip()

                st.download_button(
                    label="📄 DOCX",
                    data=create_docx(export_text),
                    file_name="lantern_draft.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )

                if HAS_REPORTLAB:
                    st.download_button(
                        label="📑 PDF",
                        data=create_pdf(export_text),
                        file_name="lantern_draft.pdf",
                        mime="application/pdf",
                        use_container_width=True
                    )
                else:
                    st.warning("Install 'reportlab' for PDF")

        st.markdown(
            f'<div class="status-pill {mode_class}" style="margin: 5px 0; width: 100%; text-align: center;">State: {mode_label}</div>',
            unsafe_allow_html=True)
        st.divider()


        # 📌 Pinned Context Section (Always Visible)
        c_pin_head, c_pin_clear = st.columns([0.8, 0.2])
        c_pin_head.markdown("<div style='font-weight: bold; font-size: 1.1em; color: #334155; margin-bottom: 10px;'>📌 Pinned Context</div>", unsafe_allow_html=True)
        if c_pin_clear.button("🗑", key="unpin_all_btn", help="Unpin all items", use_container_width=True):
            st.session_state.tree["pinned_items"] = []
            st.rerun()
        
        if st.session_state.tree["pinned_items"]:
            for i, item_data in enumerate(st.session_state.tree["pinned_items"]):
                is_dict = isinstance(item_data, dict)
                text = item_data.get("text", "") if is_dict else item_data
                title = item_data.get("title", "") if is_dict else ""
                scope = item_data.get("scope", "WD") if is_dict else "WD"
                source_context = item_data.get("source_context", "") if is_dict else ""
                
                with st.expander(f"📌 {title or 'Pinned Insight'} ({scope})"):
                    st.markdown(text)
                    # Filter out "undefined" or null contexts
                    if st.button("❌ Unpin", key=f"unpin_{i}", use_container_width=True):
                        st.session_state.tree["pinned_items"].pop(i)
                        st.rerun()
        else:
             st.markdown(
                """
                <div style="
                    background-color: #f8fafc; 
                    border: 1px solid #e2e8f0; 
                    border-radius: 6px; 
                    padding: 15px; 
                    text-align: center; 
                    color: #94a3b8; 
                    font-size: 0.9rem;">
                    No items pinned yet. Use "Pin" on suggestions to save context here.
                </div>
                """, 
                unsafe_allow_html=True
            )

        st.divider()


        # --- NEW: Scrollable Container for Suggestions ---
        with st.container(height=600):
            # 💡 Critique logic (FIXED: Tooltip + Icon + Format)
            if "current_critiques" in st.session_state and st.session_state["current_critiques"]:
                c_head, c_clear = st.columns([0.8, 0.2])
                c_head.subheader("💡 Critical Perspective")
                if c_clear.button("🗑", key="clear_all_critiques", help="Clear All Critiques", use_container_width=True):
                    st.session_state["current_critiques"] = []
                    st.rerun()

                for i, item_data in enumerate(list(st.session_state["current_critiques"])):
                    if isinstance(item_data, dict):
                        title = item_data.get("title", "Critique")
                        module = item_data.get("module", "Review")
                        text = item_data.get("text", "")
                    else:
                        title = "Critique"
                        module = "Review"
                        text = item_data

                    with st.container(border=True):
                        # Display scope and Module info
                        scope = item_data.get("scope", "Whole Document") if isinstance(item_data, dict) else "Whole Document"
                        st.markdown(
                             f'<b>{title}</b> '
                             f'<span title="Academic Factor: {module}&#10;Focus: {scope}" style="cursor: help; background-color: #38bdf8; color: white; border-radius: 4px; padding: 1px 8px; font-size: 0.7em; font-weight: bold;">i</span> '
                             f'<span style="background-color: #f1f5f9; color: #64748b; font-size: 0.75rem; padding: 2px 6px; border-radius: 4px; margin-left: 8px;">Focus: {scope}</span>',
                             unsafe_allow_html=True
                        )

                        st.markdown(
                            f'''
                            <div class="scrollable-content">
                                {text}
                            </div>
                            ''',
                            unsafe_allow_html=True
                        )

                        c_sel, c_pin, c_del = st.columns([1, 1, 1])
                        
                        with c_sel:
                             if st.button("👍", key=f"cs_sel_{i}", help="Turn critique into confidence. Mark this point as resolved to visualize the growing robustness of your argument in the dashboard", use_container_width=True):
                                # Add to persistent history (Increments Strengthened counter)
                                unique_key = text[:50]
                                if "bulletproof_history" not in st.session_state:
                                    st.session_state.bulletproof_history = set()
                                st.session_state.bulletproof_history.add(unique_key)
                                
                                # Also Pin it
                                st.session_state.tree["pinned_items"].append({
                                    "id": None,
                                    "title": title,
                                    "text": text,
                                    "type": "critique"
                                })
                                from tree import save_tree
                                save_tree(st.session_state.tree)
                                st.session_state["current_critiques"].pop(i)
                                st.rerun()

                        with c_pin:
                            if st.button("📌", key=f"cs_pin_{i}", help="Pin to context without counting as strengthened", use_container_width=True):
                                st.session_state.tree["pinned_items"].append({
                                    "id": None,
                                    "title": title,
                                    "text": text,
                                    "type": "critique"
                                })
                                save_tree(st.session_state.tree)
                                # User requested NOT to remove it from list when just pinning
                                st.rerun()

                        with c_del:
                            if st.button("🗑", key=f"cs_del_{i}", help="Delete this critique", use_container_width=True):
                                st.session_state["current_critiques"].pop(i)
                                st.rerun()
                                
            # 🌿 Suggested Paths
            if "dismissed_suggestions" not in st.session_state:
                st.session_state.dismissed_suggestions = set()

            visible_children = []
            for cid in current_node["children"]:
                if cid not in st.session_state.banned_ideas and \
                   cid not in st.session_state.dismissed_suggestions:
                    visible_children.append({"id": cid, "source": "Current"})

            if visible_children:
                c_head, c_clear = st.columns([0.8, 0.2])
                c_head.subheader("Suggested Paths")
                if c_clear.button("🗑", key="clear_all_suggestions", help="Clear All Suggestions", use_container_width=True):
                    for child_id in [item["id"] for item in visible_children]:
                        st.session_state.dismissed_suggestions.add(child_id)
                    save_tree(st.session_state.tree)
                    st.rerun()

                st.caption("Lantern generated alternative reasoning paths. Select one to continue.")
                for item in visible_children:
                    cid = item["id"]
                    child = tree["nodes"][cid]
                    meta = child.get("metadata", {})

                    title = meta.get("label")
                    explanation = meta.get("explanation")
                    module_tag = meta.get("module", "Logic")

                    raw_text = child["summary"]
                    if not title or not explanation:
                        if "Title:" in raw_text:
                            parts = re.split(r"(Title:|Module:|Explanation:|Critique:)", raw_text)
                            try:
                                if "Title:" in parts:
                                    title = parts[parts.index("Title:") + 1].strip().split('\n')[0].strip(" *")
                                if "Module:" in parts:
                                    module_tag = parts[parts.index("Module:") + 1].strip().split('\n')[0]
                                if "Explanation:" in parts:
                                    explanation = parts[parts.index("Explanation:") + 1].strip()
                                else:
                                    explanation = raw_text
                            except:
                                title = "Alternative Path"
                                explanation = raw_text
                        else:
                            title = "Alternative Path"
                            explanation = raw_text

                    if explanation:
                        explanation = re.sub(r'Title:.*?\n', '', explanation, flags=re.IGNORECASE)
                        explanation = re.sub(r'Module:.*?\n', '', explanation, flags=re.IGNORECASE)
                        explanation = explanation.replace("Title:", "").replace("Module:", "").replace("Explanation:",
                                                                                                       "").strip()

                    with st.container(border=True):
                        scope = child.get("metadata", {}).get("scope", "Whole Document")
                        st.markdown(
                            f'<div class="suggestion-meta"><span>🤖 {child.get("type", "Idea")}</span></div>',
                            unsafe_allow_html=True)

                        st.markdown(
                             f'<b>{title}</b> '
                             f'<span title="Academic Factor: {module_tag}&#10;Focus: {scope}" style="cursor: help; background-color: #38bdf8; color: white; border-radius: 4px; padding: 1px 8px; font-size: 0.7em; font-weight: bold;">i</span> '
                             f'<span style="background-color: #f1f5f9; color: #64748b; font-size: 0.75rem; padding: 2px 6px; border-radius: 4px; margin-left: 8px;">Focus: {scope}</span>',
                             unsafe_allow_html=True
                        )

                        st.markdown(
                            f'''
                            <div class="scrollable-content">
                                {explanation}
                            </div>
                            ''',
                            unsafe_allow_html=True
                        )

                        st.divider()
                        c_sel, c_pin, c_pru = st.columns([1, 1, 1])
                        with c_sel:
                            if st.button("✔", key=f"s_{cid}", help="Select this idea as your main context and active path in the Thought Tree",
                                         use_container_width=True):
                                add_debug_log(f"🖱️ SELECT: Node {cid} selected. Current text length: {len(st.session_state.get('editor_html', ''))}")
                                pin_obj = {
                                    "id": cid, 
                                    "title": title, 
                                    "text": explanation, 
                                    "type": "idea",
                                    "scope": scope,
                                    "source_context": "" # OPTIMIZATION: Do not save heavy context
                                }
                                st.session_state.tree["pinned_items"].append(pin_obj)

                                child.setdefault("metadata", {})["label"] = title
                                child.setdefault("metadata", {})["explanation"] = explanation
                                child.setdefault("metadata", {})["html"] = st.session_state.get("editor_html", "")
                                child["metadata"]["selected_path"] = True

                                # --- Automatic Sibling Dismissal ---
                                for sibling_id in current_node.get("children", []):
                                    if sibling_id != cid:
                                        st.session_state.dismissed_suggestions.add(sibling_id)

                                # --- Navigation & State Sync ---
                                if cid in st.session_state.tree.get("nodes", {}):
                                    navigate_to_node(st.session_state.tree, cid)
                                    # Use the nearest HTML (which fallback to current_node if already set above)
                                    final_html = get_nearest_html(st.session_state.tree, cid)
                                    st.session_state["editor_html"] = final_html
                                    
                                    save_tree(st.session_state.tree)
                                    
                                    # Rule 8 Exception: Navigation performance override
                                    st.rerun()
                        with c_pin:
                            if st.button("📌", key=f"p_{cid}", help="Pin this suggestion to the sidebar for future reference", use_container_width=True):
                                st.session_state.tree["pinned_items"].append({
                                    "id": cid, 
                                    "title": title, 
                                    "text": explanation, 
                                    "type": "idea",
                                    "scope": scope,
                                    "source_context": ""
                                })
                                save_tree(st.session_state.tree)
                        with c_pru:
                            if st.button("🗑", key=f"pr_{cid}", help="Dismiss this suggestion from view", use_container_width=True):
                                st.session_state.dismissed_suggestions.add(cid)
                                save_tree(st.session_state.tree)


    # ==========================================
    # FINAL AI EXECUTION LAYER (At the very bottom)
    # ==========================================
    if st.session_state.get("pending_action") and st.session_state.get("is_thinking"):
        if not st.session_state.get("llm_in_flight", False):
            payload = st.session_state.pending_action
            st.session_state.pending_action = None
            st.session_state.llm_in_flight = True
            
            try:
                # OPTIMIZATION: Process context only when absolutely necessary
                current_html = st.session_state.get("editor_html", "")
                f_mode = st.session_state.get("promo_focus_mode", "Whole Document")
                paras = st.session_state.get("structural_segments", [])
                
                # Dynamic index check for handle_event
                b_idx_raw = st.session_state.get("promo_block_selector_idx", 0)
                b_idx = b_idx_raw + 1
                
                # Use exactly what was shown in the AI Focus Preview
                t_text = st.session_state.get("focused_text", "")
                
                # Fallback only if empty (should not happen if UI rendered correctly)
                if not t_text:
                    if f_mode == "Specific Paragraph" and paras:
                        t_text = paras[max(0, min(b_idx_raw, len(paras)-1))]
                    else:
                        no_css = re.sub(r"<style.*?>.*?</style>", "", current_html, flags=re.DOTALL | re.IGNORECASE)
                        txt_s = re.sub(r"<(p|div|h[1-6]|li|blockquote|br)[^>]*>", "\n", no_css)
                        txt_s = re.sub(r"</(p|div|h[1-6]|li|blockquote)>", "\n", txt_s)
                        t_text = re.sub("<[^<]+?>", "", txt_s).replace("&nbsp;", " ").strip()
                        t_text = re.sub(r"\n{3,}", "\n\n", t_text)
                        if not t_text and current_html.strip():
                            t_text = re.sub("<[^<]+?>", "", current_html).strip()

                logger.info(f"⚡ FINAL EXEC: {payload['action'].name} | Focus={f_mode}")
                
                from controller import handle_event
                response = handle_event(st.session_state.tree, UserEventType.ACTION, {
                    "action": payload["action"],
                    "anchor_id": payload.get("anchor_id"),
                    "pinned_context": st.session_state.tree["pinned_items"],
                    "banned_ideas": st.session_state.banned_ideas,
                    "user_text": t_text,
                    "knowledge_base": st.session_state.get("knowledge_base", {}),
                    "focus_context": {"mode": f_mode, "block_idx": b_idx},
                    "logical_paragraphs": paras
                })
                
                if payload["action"] == ActionType.CRITIQUE:
                    st.session_state["current_critiques"] = response.get("items", [])
                    if not st.session_state["current_critiques"]:
                        st.session_state["ai_info_message"] = "🛡️ Lantern analyzed your draft and found it sound."
                elif payload["action"] == ActionType.DIVERGE:
                    # Suggestions appear automatically in the tree if Diverge is handled there
                    if not response.get("options"):
                        st.session_state["ai_info_message"] = "🌱 Lantern concludes the current reasoning is comprehensive."
                elif payload["action"] == ActionType.REFINE:
                    if response.get("mode") == "refine_suggestions":
                        st.session_state.pending_refine_edits = response.get("items", [])
                    else:
                        st.session_state.pending_refine_edits = [{
                            "id": f"full_refine_{os.urandom(2).hex()}",
                            "original": t_text,
                            "proposed": response.get("refined_text", ""),
                            "type": "Full Revision",
                            "reason": "Lantern provided a comprehensive revision.",
                            "status": "pending",
                            "scope": f_mode
                        }]
                elif payload["action"] == ActionType.SEGMENT:
                    # Perform structural analyzer with indicator
                    _, paragraphs = get_document_structure(current_html)
                    st.session_state.structural_segments = paragraphs
                    # SUCCESS: User requested no info message here, just rely on indicator clearing
            except Exception as e:
                st.error(f"❌ Gemini Error: {e}")
            finally:
                st.session_state.llm_in_flight = False
                st.session_state.is_thinking = False
                st.rerun()

if __name__ == "__main__":
    main()
